import os
from docx import Document
from summarize_from_txt import summarize_txt_file

def process_legal_document(docx_file_path: str):
    """
    Complete workflow: Extract text from DOCX → Generate AI Summary
    """
    print(f"🔄 Processing: {docx_file_path}")
    
    # Step 1: Extract text (your existing code)
    print("📄 Extracting text from DOCX...")
    doc = Document(docx_file_path)
    
    # Create output filename
    base_name = os.path.splitext(os.path.basename(docx_file_path))[0]
    extracted_txt_file = f"{base_name}_extracted.txt"
    
    # Extract text to file
    with open(extracted_txt_file, "w", encoding="utf-8") as f:
        f.write("--- Extracted Text ---\n\n")
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                f.write(para.text + "\n")
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    f.write(" | ".join(row_text) + "\n")
    
    print(f"✅ Text extracted to: {extracted_txt_file}")
    
    # Step 2: Generate AI summary
    print("🤖 Generating AI summary...")
    summary_result = summarize_txt_file(extracted_txt_file, save_summary=True)
    
    print(f"\n🎯 FINAL RESULTS:")
    print(f"📄 Original Document: {docx_file_path}")
    print(f"📝 Extracted Text: {extracted_txt_file}")
    print(f"🤖 AI Summary: {summary_result['summary']}")
    print(f"📊 Compression: {summary_result['stats']['compression_ratio']:.2f}")
    
    return summary_result

if __name__ == "__main__":
    # Test with your legal document
    test_docx = "test.docx"  # Your test file
    
    if os.path.exists(test_docx):
        result = process_legal_document(test_docx)
    else:
        print(f"❌ Test file '{test_docx}' not found")
