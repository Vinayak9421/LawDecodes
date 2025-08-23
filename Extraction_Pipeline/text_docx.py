from docx import Document

# Load the .docx file
doc = Document("test.docx")

# File to save extracted text
output_file = "extracted_docx.txt"

with open(output_file, "w", encoding="utf-8") as f:
    f.write("--- Extracted Text ---\n\n")

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():  # skip empty lines
            f.write(para.text + "\n")

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                f.write(" | ".join(row_text) + "\n")

print(f"\n✅ Text extracted and saved to {output_file}")
